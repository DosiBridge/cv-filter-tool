import io
from typing import Optional
import PyPDF2
from docx import Document
import aiofiles
import os

class FileProcessor:
    """Service for extracting text from PDF and Word documents"""
    
    @staticmethod
    async def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file"""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                file_data = await file.read()
            
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_data))
            text = ""
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
    
    @staticmethod
    async def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document (.docx format)"""
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"
            
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from DOCX: {str(e)}. Note: Legacy .doc files are not supported. Please convert to .docx format.")
    
    @staticmethod
    async def extract_text(file_path: str, file_extension: str) -> str:
        """Extract text from file based on extension"""
        extension = file_extension.lower()
        
        if extension == '.pdf':
            return await FileProcessor.extract_text_from_pdf(file_path)
        elif extension == '.docx':
            return await FileProcessor.extract_text_from_docx(file_path)
        elif extension == '.doc':
            # Legacy .doc format - python-docx doesn't support it
            # Attempt to open anyway, but will likely fail with helpful error
            raise ValueError(
                "Legacy .doc format is not supported. Please convert the file to .docx format. "
                "You can open the .doc file in Microsoft Word and save it as .docx."
            )
        else:
            raise ValueError(f"Unsupported file format: {extension}. Supported formats: .pdf, .docx")

